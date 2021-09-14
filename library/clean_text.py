import os
import fnmatch
import re
import spacy
import collections
import pandas as pd
from openpyxl import load_workbook


import docx

from docsim.library import dictionary_tools


def extract_paragraphs(doc_file: str):
    """Extracts paragraphs (new lines) from .docx file

    Args:
        doc_file (str): path to docx file

    Returns:
        [list]: list of text in paragraphs
    """
    doc = docx.Document(doc_file)

    paragraphs = [para.text for para in doc.paragraphs if len(para.text) > 0]

    return paragraphs


def extract_data_from_go_transcript(turns_of_talk: list):
    """Extract speaker, time, and text from turns of talk (word doc paragraphs)
     in go transcript formatting.

    Args:
        turns_of_talk (list): list of paragraphs in go transcript word doc

    Returns:
        named tuple: time_stamps, speaker_tags, text
    """
    time_stamp_regex = re.compile(r"\[[0-9:]*\]")
    time_stamps = [time_stamp_regex.findall(text) for text in turns_of_talk]
    time_stamps = [tag[0] if tag else "" for tag in time_stamps]

    turns_of_talk = [re.sub(time_stamp_regex, "", text) for text in turns_of_talk]

    speaker_tag_regex = re.compile(r"(\S[a-zA-Z1-3]+)\:")
    speaker_tags = [speaker_tag_regex.findall(text) for text in turns_of_talk]

    turns_of_talk = [re.sub(speaker_tag_regex, "", text) for text in turns_of_talk]

    if speaker_tags[2] == []:
        speaker_tag_regex = re.compile(r"([A-z]+\s[A-z]+[\s1-9]+)\:")
        speaker_tags = [speaker_tag_regex.findall(text) for text in turns_of_talk]
        turns_of_talk = [re.sub(speaker_tag_regex, "", text) for text in turns_of_talk]

    if speaker_tags[2] == []:
        speaker_tag_regex = re.compile(r"([A-z]+[\s1-9]+)\:")
        speaker_tags = [speaker_tag_regex.findall(text) for text in turns_of_talk]
        turns_of_talk = [re.sub(speaker_tag_regex, "", text) for text in turns_of_talk]

    Transcript = collections.namedtuple(
        "Transcript", ["time_stamps", "speaker_tags", "text"]
    )

    speaker_tags = [tag[0].strip() if tag else "" for tag in speaker_tags]

    transcript = Transcript(
        time_stamps=time_stamps, speaker_tags=speaker_tags, text=turns_of_talk
    )

    return transcript


def extract_data_from_otter_transcript(turns_of_talk: list):
    """Extract speaker, time, and text from turns of talk (word doc paragraphs)
     in otter transcript formatting.

    Args:
        turns_of_talk (list): list of paragraphs in otter transcript word doc

    Returns:
        named tuple: time_stamps, speaker_tags, text
    """
    speaker_tags = []
    time_stamps = []
    texts = []
    turn = 6
    while turn < len(turns_of_talk):
        speaker_tags.append(turns_of_talk[turn][:-5].strip())
        time_stamps.append(turns_of_talk[turn][-5:])
        turn = turn + 1

        texts.append(turns_of_talk[turn])
        turn = turn + 1

    Transcript = collections.namedtuple(
        "Transcript", ["time_stamps", "speaker_tags", "text"]
    )

    transcript = Transcript(
        time_stamps=time_stamps, speaker_tags=speaker_tags, text=texts
    )
    return transcript


def word_to_transcript(doc_file=str):
    """Extracts paragraphs from word doc and directs
    to appropriate processing function

    Args:
        doc_file ([str], optional): String path to file

    Returns:
        named tuple: time_stamps, speaker_tags, text
    """
    paragraphs = extract_paragraphs(doc_file=doc_file)
    if "[0" in paragraphs[0]:
        return extract_data_from_go_transcript(turns_of_talk=paragraphs)
    if paragraphs[2] == "SUMMARY KEYWORDS":
        return extract_data_from_otter_transcript(turns_of_talk=paragraphs)

    else:
        print("error with " + doc_file)


def import_text(filepath: str, pattern: str, paragraph_tag: str = None):
    """Searches folder for docx files and extracts plain text with a tag indicating each new paragraph

    Args:
        filepath (str): filepath to folder containing docx files
        pattern (str): search pattern

    Returns:
        [dict]: dictionary of filenames and text
    """
    doc_dict = {}
    for filename in os.listdir(filepath):
        if fnmatch.fnmatch(filename, pattern) and not filename.startswith("~$"):
            print(filename)
            docfile = filepath + filename
            doc = docx.Document(docfile)
            text_list = []
            for para in doc.paragraphs:
                if paragraph_tag is None:
                    text_list.append(para.text)
                elif paragraph_tag and (len(para.text) > 0):
                    segmented_text = paragraph_tag + para.text
                    text_list.append(segmented_text)
            text = " ".join(text_list)
            doc_dict[filename] = text

    return doc_dict


def filter_segments(
    text: str, segment_tag: str, filter_tag: str, keep_if_tagged: bool = True
):
    """Filters a text by searching individual segments for instance of string

    Args:
        text (str): string to filter
        segment_tag (str): tag indicating text segments to search
        filter_tag (str): string to search for
        keep_if_tagged (bool, optional): Keep segment if contains filter_tag? Defaults to True.

    Returns:
        [str]: filtered string
    """
    segments = [
        segment_tag + " " + string.strip()
        for string in text.split(segment_tag)
        if string
    ]
    filtered_segments = [segment for segment in segments if filter_tag in segment]
    filtered_text = " ".join(filtered_segments)

    return filtered_text


def add_whitespace_after_punct(s):
    return re.sub(r"([.?!])", ". ", re.sub(r" +", " ", s))


def remove_trailing_hyphen(s):
    return re.sub(r"(-+)\s", " ", re.sub(r" +", " ", s))


def remove_leading_hyphen(s):
    return re.sub(r"\s(-+)", " ", re.sub(r" +", " ", s))


def remove_hyphens(s):
    return s.replace("-", " ")


def word_family_from_dict(text: str, families: dict):
    """Replace words found in dictionary values (list) with key

    Args:
        text (str): String with substrings to replace
        families (dict): Word family name in key, list of members as value

    Returns:
        new_text: String with word family members replaced with word family name
    """
    new_dict = dictionary_tools.long_inverse_dict_from_key_list(families)

    old_words = list(new_dict.keys())

    new_text = text
    for old_word in old_words:
        new_text = new_text.replace(old_word, new_dict[old_word])
        new_text = new_text.replace(
            " " + old_word + " ", " " + new_dict[old_word] + " "
        )
        new_text = new_text.replace(
            " " + old_word + ".", " " + new_dict[old_word] + "."
        )
        new_text = new_text.replace(
            " " + old_word + ",", " " + new_dict[old_word] + ","
        )
        new_text = new_text.replace(
            "[" + old_word + " ", " " + "[" + new_dict[old_word] + " "
        )

    # for word in new_dict:
    #     new_text = new_text.replace(word, new_dict[word])

    return new_text


def transcript_to_cleaned_df(transcript):
    """Turns transcript named tuple to dataframe

    Args:
        transcript ([named tuple]): transcript named tuple

    Returns:
        [df]: dataframe with three columns: time, speaker, text
    """
    new_speaker_tags = []
    current_speaker = ""
    for speaker in transcript.speaker_tags:
        if speaker:
            current_speaker = speaker
            new_speaker_tags.append(speaker)
        else:
            new_speaker_tags.append(current_speaker)

    new_time_tags = []
    current_time = ""
    for time in transcript.time_stamps:
        if time:
            current_time = time
            new_time_tags.append(current_time)
        else:
            new_time_tags.append(current_time)

    df = pd.DataFrame(
        {
            "time": new_time_tags,
            "speaker": new_speaker_tags,
            "text": transcript.text,
        }
    )

    df["full_text"] = df.groupby(["speaker", "time"])["text"].transform(
        lambda x: " ".join(x)
    )
    df = df.drop_duplicates(subset=["speaker", "time", "full_text"], keep="first")
    df = df[["speaker", "time", "full_text"]].rename(columns={"full_text": "text"})
    return df


def transcript_df_to_excel(transcript_df, excel_template: str, excel_file: str):

    wb = load_workbook(excel_template)
    ws = wb.active

    row = 2
    col = 1
    for time in transcript_df.time:
        ws.cell(row=row, column=col).value = str(time)
        row = row + 1

    row = 2
    col = 2

    for speaker in transcript_df.speaker:
        ws.cell(row=row, column=col).value = str(speaker)
        row = row + 1

    row = 2
    col = 3
    for text in transcript_df.text:
        ws.cell(row=row, column=col).value = text
        row = row + 1

    wb.save(excel_file)
